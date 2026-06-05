import docx
import sys
import os

def update_document(input_path, output_path):
    print(f"Reading {input_path}")
    doc = docx.Document(input_path)
    in_model_section = False
    
    for i, para in enumerate(doc.paragraphs):
        # 1. Update Database tech text
        if "como Django e MySQL" in para.text:
            para.text = para.text.replace("como Django e MySQL", "como Django, SQLite (para testes/desenvolvimento) e PostgreSQL (para produção)")
            
        if "como Django, SQLite e PostgreSQL" in para.text:
            para.text = para.text.replace("como Django, SQLite e PostgreSQL", "como Django, SQLite (para testes/desenvolvimento) e PostgreSQL (para produção)")

        # 2. Update DB entities 
        if "10.3. MODELO CONCEITUAL DO BANCO DE DADOS" in para.text:
            in_model_section = True
            
        if in_model_section:
            if "Empresa: Armazena informações da empresa terceirizada" in para.text:
                para.text = "Perfil de Usuário (UserProfile): Contém dados dos usuários, seus respectivos perfis (Admin, Gestor, Líder, Colaborador) e as informações integradas da empresa terceirizada (nome, CNPJ, logradouro, etc.)."
                
            if "Usuário: Contém dados dos usuários" in para.text:
                para.text = "Especialidade: Categoriza as diferentes áreas de atuação dos profissionais técnicos."
                
    # Second pass for insertions
    for i, para in enumerate(doc.paragraphs):
        # Insert new DB entities
        if "Mensagem: Armazena o histórico de comunicação de cada tarefa." in para.text:
            # check if they were already added to prevent duplicates
            next_p = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
            if next_p and "Evidência de Tarefa" not in next_p.text:
                next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else doc.paragraphs[i]
                next_para.insert_paragraph_before("Evidência de Tarefa (TaskEvidence): Registra fotos e descrições para validação da conclusão.")
                next_para.insert_paragraph_before("Notificação: Gerencia os alertas do sistema sobre o andamento das tarefas.")
                next_para.insert_paragraph_before("Log de Atividades (ActivityLog): Mantém registro de auditoria sobre ações críticas no sistema.")

        # Remove 2FA mentions if they were added previously
        if "Autenticação de Dois Fatores (2FA)" in para.text:
            p = para._element
            p.getparent().remove(p)
            para._p = para._element = None

        if "11. CONSIDERAÇÕES FINAIS" in para.text:
            # Insert section if it doesn't exist
            prev_p = doc.paragraphs[i-1] if i > 0 else None
            if prev_p and "RECURSOS ADICIONAIS" not in prev_p.text and prev_p.text.strip() == "":
                 pass # Might need more complex logic to check for existence, let's assume we use the original file.

    # Instead of complicated duplicate checking, let's just use the original docx as base
    # But wait, to be safe, if we use the original doc, we know exactly what is there.
    pass

def process_original():
    input_path = "TCC_CHAPLIN-Corrigido_19-05.docx"
    output_path = "TCC_CHAPLIN-Corrigido_19-05_Final.docx"
    doc = docx.Document(input_path)
    in_model_section = False
    
    for para in doc.paragraphs:
        if "como Django e MySQL" in para.text:
            para.text = para.text.replace("como Django e MySQL", "como Django, SQLite (para testes/desenvolvimento) e PostgreSQL (para produção)")
            
        if "10.3. MODELO CONCEITUAL DO BANCO DE DADOS" in para.text:
            in_model_section = True
            
        if in_model_section:
            if "Empresa: Armazena informações da empresa terceirizada" in para.text:
                para.text = "Perfil de Usuário (UserProfile): Contém dados dos usuários, seus respectivos perfis (Admin, Gestor, Líder, Colaborador) e as informações integradas da empresa terceirizada (nome, CNPJ, logradouro, etc.)."
                
            if "Usuário: Contém dados dos usuários" in para.text:
                para.text = "Especialidade: Categoriza as diferentes áreas de atuação dos profissionais técnicos."
                
    for i, para in enumerate(doc.paragraphs):
        if "Mensagem: Armazena o histórico de comunicação de cada tarefa." in para.text:
            next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else doc.paragraphs[i]
            next_para.insert_paragraph_before("Evidência de Tarefa (TaskEvidence): Registra fotos e descrições para validação da conclusão.")
            next_para.insert_paragraph_before("Notificação: Gerencia os alertas do sistema sobre o andamento das tarefas.")
            next_para.insert_paragraph_before("Log de Atividades (ActivityLog): Mantém registro de auditoria sobre ações críticas no sistema.")

        if "11. CONSIDERAÇÕES FINAIS" in para.text:
            new_title = doc.paragraphs[i].insert_paragraph_before("10.4. RECURSOS ADICIONAIS IMPLEMENTADOS")
            new_title.style = doc.paragraphs[i].style
            doc.paragraphs[i].insert_paragraph_before("Além dos requisitos básicos, o sistema conta com recursos técnicos adicionais que garantem maior segurança e usabilidade:")
            doc.paragraphs[i].insert_paragraph_before("- Integração com ViaCEP: Permite a busca automática e a padronização dos endereços tanto para as tarefas de manutenção quanto para os perfis de usuários.")
            doc.paragraphs[i].insert_paragraph_before("")

    doc.save(output_path)
    
    # Overwrite the Atualizado one as well to avoid confusion
    doc.save("TCC_CHAPLIN-Corrigido_19-05_Atualizado.docx")
    print("Documents successfully updated to Final.")

if __name__ == "__main__":
    process_original()
