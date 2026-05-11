import docx
import sys

def update_document(file_path):
    doc = docx.Document(file_path)

    in_model_section = False
    
    for para in doc.paragraphs:
        if "como Django e MySQL" in para.text:
            para.text = para.text.replace("como Django e MySQL", "como Django, SQLite e PostgreSQL")
            
        if "10.3. MODELO CONCEITUAL DO BANCO DE DADOS" in para.text:
            in_model_section = True
            
        if in_model_section:
            if "Empresa: Armazena informações da empresa terceirizada" in para.text:
                para.text = "Perfil de Usuário (UserProfile): Contém dados dos usuários, seus respectivos perfis (Admin, Gestor, Líder, Colaborador) e as informações integradas da empresa terceirizada (nome, CNPJ, logradouro, etc.)."
                
            if "Usuário: Contém dados dos usuários" in para.text:
                para.text = "Especialidade: Categoriza as diferentes áreas de atuação dos profissionais técnicos."
                
            if "Mensagem: Armazena o histórico" in para.text:
                pass 

    # Second pass for inserting new entities and new section
    for i, para in enumerate(doc.paragraphs):
        if "Mensagem: Armazena o histórico de comunicação de cada tarefa." in para.text:
            next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else doc.paragraphs[i]
            next_para.insert_paragraph_before("Evidência de Tarefa (TaskEvidence): Registra fotos e descrições para validação da conclusão.")
            next_para.insert_paragraph_before("Notificação: Gerencia os alertas do sistema sobre o andamento das tarefas.")
            next_para.insert_paragraph_before("Log de Atividades (ActivityLog): Mantém registro de auditoria sobre ações críticas no sistema.")

        if "11. CONSIDERAÇÕES FINAIS" in para.text:
            # Insert a new section before Considerações Finais
            new_title = doc.paragraphs[i].insert_paragraph_before("10.4. RECURSOS ADICIONAIS IMPLEMENTADOS")
            new_title.style = doc.paragraphs[i].style # Try to match heading style
            doc.paragraphs[i].insert_paragraph_before("Além dos requisitos básicos, o sistema conta com recursos técnicos adicionais que garantem maior segurança e usabilidade:")
            doc.paragraphs[i].insert_paragraph_before("- Autenticação de Dois Fatores (2FA) e Verificação de E-mail: Implementados no perfil do usuário para fornecer uma camada extra de segurança.")
            doc.paragraphs[i].insert_paragraph_before("- Integração com ViaCEP: Permite a busca automática e a padronização dos endereços tanto para as tarefas de manutenção quanto para os perfis de usuários.")
            doc.paragraphs[i].insert_paragraph_before("")

    doc.save("TCC_CHAPLIN-Corrigido_05-05_Atualizado.docx")
    print("Document successfully updated.")

if __name__ == "__main__":
    update_document(sys.argv[1])
